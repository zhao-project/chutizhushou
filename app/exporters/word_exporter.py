"""试卷导出器 - 生成 Word/PDF"""

import io
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from typing import List
from app.models.schemas import ExamPaper, Question, QuestionType


def _format_question_number(idx: int) -> str:
    return f"{idx + 1}"


def _write_question(para, q: Question, idx: int, include_answer: bool = False):
    """写一道题目到段落"""
    type_labels = {
        QuestionType.CHOICE: "选择题",
        QuestionType.FILL_BLANK: "填空题",
        QuestionType.TRUE_FALSE: "判断题",
        QuestionType.SHORT_ANSWER: "简答题",
        QuestionType.ESSAY: "论述题",
    }
    label = type_labels.get(q.question_type, "")

    # 题号 + 题型 + 题干
    run = para.add_run(f"{_format_question_number(idx)}. 【{label}】{q.content}")
    run.font.size = Pt(12)
    run.font.name = "宋体"

    # 选项
    if q.options:
        para.add_run("\n")
        for opt in q.options:
            run_opt = para.add_run(f"    {opt}\n")
            run_opt.font.size = Pt(11)
            run_opt.font.name = "宋体"

    # 答案（答案分离模式下不在试卷中显示）
    if include_answer and q.answer:
        para.add_run("\n")
        run_ans = para.add_run(f"    答案：{q.answer}")
        run_ans.font.size = Pt(11)
        run_ans.font.color.rgb = RGBColor(0, 0, 128)
        if q.explanation:
            para.add_run("\n")
            run_exp = para.add_run(f"    解析：{q.explanation}")
            run_exp.font.size = Pt(10)
            run_exp.font.color.rgb = RGBColor(128, 128, 128)


def export_word(paper: ExamPaper, separate_answers: bool = True, include_answer_sheet: bool = False) -> bytes:
    """导出为 Word (.docx) 格式，返回字节"""
    doc = Document()

    # 页面设置
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    # === 试卷标题 ===
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title.add_run(paper.title)
    run_title.font.size = Pt(18)
    run_title.bold = True
    run_title.font.name = "黑体"

    # 副标题信息
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_text = f"科目：{paper.subject}"
    if paper.grade:
        info_text += f"  |  年级：{paper.grade}"
    if paper.chapter:
        info_text += f"  |  章节：{paper.chapter}"
    info_text += f"  |  总分：{paper.total_score}分  |  时长：{paper.duration_min}分钟"
    run_info = info.add_run(info_text)
    run_info.font.size = Pt(11)
    run_info.font.name = "宋体"

    doc.add_paragraph("")  # 空行

    # === 题目 ===
    current_type = None
    for idx, q in enumerate(paper.questions):
        # 按题型分组
        type_labels = {
            QuestionType.CHOICE: "一、选择题",
            QuestionType.FILL_BLANK: "二、填空题",
            QuestionType.TRUE_FALSE: "三、判断题",
            QuestionType.SHORT_ANSWER: "四、简答题",
            QuestionType.ESSAY: "五、论述题",
        }
        type_header = type_labels.get(q.question_type, "")
        if q.question_type != current_type and type_header:
            current_type = q.question_type
            header = doc.add_paragraph()
            run_h = header.add_run(type_header)
            run_h.font.size = Pt(14)
            run_h.bold = True
            run_h.font.name = "黑体"

        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(8)
        _write_question(para, q, idx, include_answer=not separate_answers)

    # === 答案部分（分离模式） ===
    if separate_answers:
        doc.add_page_break()
        answer_title = doc.add_paragraph()
        answer_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_at = answer_title.add_run("参考答案与解析")
        run_at.font.size = Pt(16)
        run_at.bold = True
        run_at.font.name = "黑体"
        doc.add_paragraph("")

        for idx, q in enumerate(paper.questions):
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(6)
            run_num = para.add_run(f"{_format_question_number(idx)}. ")
            run_num.bold = True
            run_num.font.size = Pt(11)
            run_ans = para.add_run(f"{q.answer}")
            run_ans.font.size = Pt(11)
            if q.explanation:
                para.add_run("\n")
                run_exp = para.add_run(f"   解析：{q.explanation}")
                run_exp.font.size = Pt(10)

    # === 答题卡（可选） ===
    if include_answer_sheet:
        doc.add_page_break()
        sheet_title = doc.add_paragraph()
        sheet_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_st = sheet_title.add_run("答 题 卡")
        run_st.font.size = Pt(16)
        run_st.bold = True
        run_st.font.name = "黑体"

        # 选择答题区域
        choice_qs = [(i, q) for i, q in enumerate(paper.questions) if q.question_type == QuestionType.CHOICE]
        if choice_qs:
            doc.add_paragraph("")
            run_ch = doc.add_paragraph().add_run("选择题答题区：")
            run_ch.bold = True
            run_ch.font.size = Pt(12)

            # 每行5题
            for row_start in range(0, len(choice_qs), 5):
                row_qs = choice_qs[row_start:row_start + 5]
                row_para = doc.add_paragraph()
                for i, q in row_qs:
                    run_q = row_para.add_run(f"{_format_question_number(i)}. ___  ")
                    run_q.font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
